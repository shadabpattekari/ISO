"""
FaizZab Toolkit Generation Engine — CORE POC (isolation test).

Proves the hardest, most failure-prone part BEFORE building the app:
  * Deterministic, branded DOCX / XLSX / PDF generation from template specs +
    client / industry / standard variables.
  * ZIP packaging matching a manifest.
  * Quality validation (no unresolved placeholders, mandatory sections, unique
    doc IDs, files openable, ZIP completeness, logo present).

Run:  python /app/poc/test_core.py
"""
import io
import json
import os
import sys
import zipfile

sys.path.insert(0, "/app/backend")

from PIL import Image, ImageDraw, ImageFont  # noqa: E402
from docx import Document as DocxReader  # noqa: E402
from openpyxl import load_workbook  # noqa: E402

from toolkit_engine import (  # noqa: E402
    render_docx_document,
    render_pdf_document,
    render_xlsx_register,
    build_zip,
    validate_document,
)

OUT = "/app/poc/output"
os.makedirs(OUT, exist_ok=True)


# --------------------------------------------------------------------------- #
# Test fixtures
# --------------------------------------------------------------------------- #
def make_logo(path: str, text: str, color: str) -> None:
    img = Image.new("RGB", (400, 400), color=color)
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 120)
    except Exception:
        font = ImageFont.load_default()
    d.ellipse((30, 30, 370, 370), outline="white", width=10)
    bbox = d.textbbox((0, 0), text, font=font)
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    d.text(((400 - w) / 2, (400 - h) / 2 - 20), text, fill="white", font=font)
    img.save(path)


def context_saas() -> dict:
    make_logo(f"{OUT}/logo_saas.png", "AC", "#2563EB")
    return {
        "org": {
            "legal_name": "AcmeCloud Technologies Private Limited",
            "trade_name": "AcmeCloud",
            "website": "acmecloud.io",
            "industry": "SaaS",
            "employee_count": "35",
            "brand_color": "#2563EB",
            "logo_path": f"{OUT}/logo_saas.png",
            "products_services": "Multi-tenant SaaS analytics platform",
            "locations": ["Bengaluru HQ", "Pune Dev Center"],
        },
        "roles": {
            "isms_manager": "Priya Sharma (CISO)",
            "internal_auditor": "Rahul Verma",
            "top_management": "Anita Desai (CEO)",
        },
        "doc_control": {
            "version": "1.0",
            "effective_date": "01 Feb 2026",
            "review_date": "01 Feb 2027",
            "classification": "Confidential",
            "prepared_by": "Priya Sharma",
            "reviewed_by": "Rahul Verma",
            "approved_by": "Anita Desai",
        },
    }


def context_healthcare() -> dict:
    make_logo(f"{OUT}/logo_health.png", "MH", "#0F766E")
    return {
        "org": {
            "legal_name": "MediHeal Diagnostics LLP",
            "trade_name": "MediHeal",
            "website": "mediheal.in",
            "industry": "Healthcare",
            "employee_count": "48",
            "brand_color": "#0F766E",
            "logo_path": f"{OUT}/logo_health.png",
            "products_services": "Diagnostic laboratory and imaging services",
            "locations": ["Chennai Main Lab", "Coimbatore Collection Center", "Madurai Center"],
        },
        "roles": {
            "isms_manager": "Dr. Suresh Kumar",
            "internal_auditor": "Lakshmi Nair",
            "top_management": "Dr. Ramesh Iyer (Managing Partner)",
        },
        "doc_control": {
            "version": "1.0",
            "effective_date": "15 Jan 2026",
            "review_date": "15 Jan 2027",
            "classification": "Restricted",
            "prepared_by": "Dr. Suresh Kumar",
            "reviewed_by": "Lakshmi Nair",
            "approved_by": "Dr. Ramesh Iyer",
        },
    }


# --------------------------------------------------------------------------- #
# Sample approved template specs (representative of real toolkit content)
# --------------------------------------------------------------------------- #
def policy_spec() -> dict:
    return {
        "doc_id": "ISMS-POL-001",
        "title": "Information Security Policy",
        "standard": "ISO/IEC 27001:2022",
        "classification": "Confidential",
        "clause_refs": ["5.2", "5.1", "A.5.1"],
        "owner_role": "{{roles.isms_manager}}",
        "sections": [
            {"heading": "Purpose", "mandatory": True,
             "body": "This Information Security Policy establishes {{org.legal_name}}'s commitment to "
                     "protecting the confidentiality, integrity and availability of information assets "
                     "used to deliver {{org.products_services}}. It applies across {{org.locations}}."},
            {"heading": "Scope", "mandatory": True,
             "body": "This policy applies to all employees, contractors and third parties of "
                     "{{org.trade_name}} (approximately {{org.employee_count}} personnel) operating "
                     "within the defined ISMS scope."},
            {"heading": "Policy Statements", "mandatory": True,
             "body": "{{org.trade_name}} shall:",
             "bullets": [
                 "Manage information security risks in line with a documented methodology.",
                 "Ensure compliance with applicable legal, regulatory and contractual obligations.",
                 "Provide security awareness training to all personnel.",
                 "Maintain, review and continually improve the ISMS.",
             ]},
            {"heading": "Roles and Responsibilities", "mandatory": True,
             "body": "Top management ({{roles.top_management}}) is accountable for the ISMS. "
                     "The ISMS Manager ({{roles.isms_manager}}) is responsible for day-to-day operation. "
                     "The Internal Auditor ({{roles.internal_auditor}}) independently audits the ISMS."},
        ],
        "responsibilities": [
            ("{{roles.top_management}}", "Overall accountability and resource provision for the ISMS."),
            ("{{roles.isms_manager}}", "Operation, monitoring and maintenance of the ISMS."),
            ("All personnel", "Compliance with security policies and reporting of incidents."),
        ],
        "kpis": [
            "Percentage of staff completing annual security awareness training (target 100%).",
            "Number of security incidents resolved within SLA (target 95%).",
        ],
        "common_nonconformities": [
            "Policy not reviewed within the defined review period.",
            "No evidence of top-management approval.",
            "Awareness training records incomplete.",
        ],
        "audit_questions": [
            "Has the information security policy been approved by top management?",
            "Is the policy communicated to all relevant personnel?",
            "Is there evidence of periodic review?",
        ],
        "related_documents": [
            "ISMS Scope (ISMS-DOC-002)",
            "Risk Assessment Methodology (ISMS-PROC-003)",
            "Statement of Applicability (ISMS-DOC-010)",
        ],
    }


def register_spec() -> dict:
    return {
        "doc_id": "ISMS-REG-004",
        "title": "Information Security Risk Register",
        "standard": "ISO/IEC 27001:2022",
        "sheet_name": "Risk Register",
        "columns": [
            {"label": "Risk ID", "key": "id", "width": 12},
            {"label": "Asset / Process", "key": "asset", "width": 26},
            {"label": "Threat", "key": "threat", "width": 28},
            {"label": "Likelihood", "key": "likelihood", "width": 14},
            {"label": "Impact", "key": "impact", "width": 12},
            {"label": "Risk Level", "key": "level", "width": 14},
            {"label": "Treatment", "key": "treatment", "width": 30},
            {"label": "Owner", "key": "owner", "width": 22},
        ],
        "rows": [
            {"id": "R-001", "asset": "Customer data ({{org.trade_name}})", "threat": "Unauthorized access",
             "likelihood": "Medium", "impact": "High", "level": "High",
             "treatment": "Access control + MFA", "owner": "{{roles.isms_manager}}"},
            {"id": "R-002", "asset": "Production infrastructure", "threat": "Service outage",
             "likelihood": "Low", "impact": "High", "level": "Medium",
             "treatment": "Backup + DR plan", "owner": "{{roles.isms_manager}}"},
        ],
    }


def dpdpa_spec() -> dict:
    return {
        "doc_id": "DPDP-POL-001",
        "title": "Data Protection Governance Policy",
        "standard": "India DPDP Act 2023 & DPDP Rules 2025",
        "classification": "Confidential",
        "clause_refs": ["S.8", "S.5"],
        "owner_role": "{{roles.isms_manager}}",
        "sections": [
            {"heading": "Purpose", "mandatory": True,
             "body": "This policy sets out how {{org.legal_name}} acts as a Data Fiduciary under the "
                     "Digital Personal Data Protection Act, 2023 when processing digital personal data."},
            {"heading": "Lawful Processing", "mandatory": True,
             "body": "{{org.trade_name}} processes personal data only for lawful purposes with valid "
                     "consent or under legitimate uses permitted by the Act."},
        ],
        "disclaimer": "This document is a template provided by FaizZab to assist {{org.legal_name}} with "
                      "DPDP Act readiness. It does NOT constitute legal advice. {{org.trade_name}} must "
                      "obtain independent legal review before relying on this document.",
    }


# --------------------------------------------------------------------------- #
# Output verification helpers
# --------------------------------------------------------------------------- #
def verify_docx(data: bytes) -> tuple:
    try:
        d = DocxReader(io.BytesIO(data))
        text = "\n".join(p.text for p in d.paragraphs)
        tables = len(d.tables)
        if "{{" in text:
            return False, "unresolved placeholder in DOCX text"
        return True, f"opened OK, {len(d.paragraphs)} paragraphs, {tables} tables"
    except Exception as e:
        return False, f"DOCX open failed: {e}"


def verify_xlsx(data: bytes) -> tuple:
    try:
        wb = load_workbook(io.BytesIO(data))
        ws = wb.active
        cells = list(ws.iter_rows(values_only=True))
        flat = " ".join(str(c) for row in cells for c in row if c)
        if "{{" in flat:
            return False, "unresolved placeholder in XLSX"
        return True, f"opened OK, sheet '{ws.title}', {ws.max_row} rows x {ws.max_column} cols"
    except Exception as e:
        return False, f"XLSX open failed: {e}"


def verify_pdf(data: bytes) -> tuple:
    if not data.startswith(b"%PDF"):
        return False, "not a valid PDF (missing %PDF header)"
    if b"%%EOF" not in data[-1024:]:
        return False, "PDF missing EOF marker"
    return True, f"valid PDF, {len(data)} bytes"


def verify_zip(data: bytes, expected: list) -> tuple:
    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
        bad = zf.testzip()
        if bad:
            return False, f"corrupt entry: {bad}"
        names = zf.namelist()
        missing = [e for e in expected if e not in names]
        if missing:
            return False, f"ZIP missing expected files: {missing}"
        return True, f"complete ZIP, {len(names)} entries"
    except Exception as e:
        return False, f"ZIP open failed: {e}"


# --------------------------------------------------------------------------- #
# Main POC run
# --------------------------------------------------------------------------- #
def run_for(profile_name: str, context: dict) -> dict:
    report = {"profile": profile_name, "checks": [], "documents": [], "ok": True}

    def check(name, ok, detail):
        report["checks"].append({"name": name, "ok": bool(ok), "detail": detail})
        if not ok:
            report["ok"] = False

    specs = {
        "ISMS-POL-001_Information_Security_Policy": ("docx", policy_spec()),
        "DPDP-POL-001_Data_Protection_Governance_Policy": ("docx", dpdpa_spec()),
        "ISMS-REG-004_Risk_Register": ("xlsx", register_spec()),
    }

    # Validation pass (pre-generation)
    doc_ids = []
    for _, (_, spec) in specs.items():
        val = validate_document(spec, context)
        report["documents"].append(val)
        check(f"validate:{spec['doc_id']}", val["ok"], val["errors"] or "no errors")
        doc_ids.append(spec["doc_id"])

    # Unique doc IDs
    check("unique_doc_ids", len(doc_ids) == len(set(doc_ids)), doc_ids)

    # Logo present
    check("logo_present", bool(context["org"].get("logo_path") and os.path.exists(context["org"]["logo_path"])),
          context["org"].get("logo_path"))

    # Generate files
    zip_files = {}
    prefix = profile_name.replace(" ", "_")
    for base, (kind, spec) in specs.items():
        if kind == "docx":
            data = render_docx_document(spec, context)
            ok, detail = verify_docx(data)
            check(f"docx:{base}", ok, detail)
            fname = f"{base}.docx"
            zip_files[f"documents/{fname}"] = data
            # PDF reference copy
            pdf = render_pdf_document(spec, context)
            pok, pdetail = verify_pdf(pdf)
            check(f"pdf:{base}", pok, pdetail)
            zip_files[f"reference_pdf/{base}.pdf"] = pdf
            with open(f"{OUT}/{prefix}_{fname}", "wb") as fh:
                fh.write(data)
            with open(f"{OUT}/{prefix}_{base}.pdf", "wb") as fh:
                fh.write(pdf)
        else:
            data = render_xlsx_register(spec, context)
            ok, detail = verify_xlsx(data)
            check(f"xlsx:{base}", ok, detail)
            fname = f"{base}.xlsx"
            zip_files[f"registers/{fname}"] = data
            with open(f"{OUT}/{prefix}_{fname}", "wb") as fh:
                fh.write(data)

    # Manifest + index (manifest lists all artifact files generated so far)
    artifact_files = list(zip_files.keys())
    manifest = {
        "organization": context["org"]["legal_name"],
        "generated_documents": artifact_files,
        "count": len(artifact_files),
    }
    zip_files["toolkit_manifest.json"] = json.dumps(manifest, indent=2).encode()
    index_html = f"<html><body><h1>{context['org']['legal_name']} — Toolkit Index</h1><ul>" + \
        "".join(f"<li>{k}</li>" for k in zip_files) + "</ul></body></html>"
    zip_files["index.html"] = index_html.encode()

    zip_bytes = build_zip(zip_files)
    expected = list(zip_files.keys())
    zok, zdetail = verify_zip(zip_bytes, expected)
    check("zip_complete", zok, zdetail)
    with open(f"{OUT}/{prefix}_toolkit.zip", "wb") as fh:
        fh.write(zip_bytes)

    # Manifest matches outputs: every file listed in the manifest exists in the ZIP
    zf_names = set(zipfile.ZipFile(io.BytesIO(zip_bytes)).namelist())
    manifest_ok = all(f in zf_names for f in manifest["generated_documents"])
    check("manifest_matches_outputs", manifest_ok,
          f"{manifest['count']} artifacts all present in ZIP")

    return report


def main():
    all_reports = []
    for name, ctx in [("SaaS", context_saas()), ("Healthcare", context_healthcare())]:
        all_reports.append(run_for(name, ctx))

    overall_ok = all(r["ok"] for r in all_reports)
    summary = {"overall_ok": overall_ok, "reports": all_reports}
    with open(f"{OUT}/generation_report.json", "w") as fh:
        json.dump(summary, fh, indent=2)

    print("=" * 70)
    print("FaizZab Toolkit Engine — CORE POC")
    print("=" * 70)
    for r in all_reports:
        print(f"\n[{r['profile']}]  overall_ok={r['ok']}")
        for c in r["checks"]:
            status = "PASS" if c["ok"] else "FAIL"
            print(f"  [{status}] {c['name']}: {c['detail']}")
    print("\n" + "=" * 70)
    print(f"OVERALL: {'PASS ✅' if overall_ok else 'FAIL ❌'}")
    print(f"Outputs written to {OUT}")
    print("=" * 70)
    sys.exit(0 if overall_ok else 1)


if __name__ == "__main__":
    main()
