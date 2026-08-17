"""
Core rendering engine for FaizZab toolkit documents.

Everything here is pure (no DB / no network). It takes a *document spec*
(the approved controlled structure of a document) plus a *context* (client,
brand, doc-control, roles, locations, industry variables) and produces bytes
for DOCX / PDF / XLSX, plus a validation report.
"""
from __future__ import annotations

import io
import re
import zipfile
from datetime import datetime
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage,
    PageBreak, KeepTogether,
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT

PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_.]+)\s*\}\}")


# --------------------------------------------------------------------------- #
# Variable resolution
# --------------------------------------------------------------------------- #
def flatten_context(context: Dict[str, Any], prefix: str = "") -> Dict[str, str]:
    """Flatten a nested dict into dot-notation keys -> string values."""
    out: Dict[str, str] = {}
    for key, val in context.items():
        full = f"{prefix}.{key}" if prefix else key
        if isinstance(val, dict):
            out.update(flatten_context(val, full))
        elif isinstance(val, (list, tuple)):
            # Join list values for text interpolation
            rendered = ", ".join(str(v) for v in val if not isinstance(v, dict))
            out[full] = rendered
        elif val is None:
            out[full] = ""
        else:
            out[full] = str(val)
    return out


def resolve_variables(text: str, flat: Dict[str, str]) -> str:
    """Replace {{ dotted.key }} tokens with values from the flat context."""
    if not text:
        return text

    def _sub(match: "re.Match[str]") -> str:
        key = match.group(1)
        return flat.get(key, match.group(0))  # keep token if missing (caught by validator)

    return PLACEHOLDER_RE.sub(_sub, text)


def find_unresolved(text: str) -> List[str]:
    return PLACEHOLDER_RE.findall(text or "")


def hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    h = (hex_color or "#1F3A5F").lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    try:
        return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    except Exception:
        return 31, 58, 95


# --------------------------------------------------------------------------- #
# DOCX rendering
# --------------------------------------------------------------------------- #
def _set_cell_background(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)


def _add_doc_control_header_footer(doc: Document, ctx_flat: Dict[str, str], spec: Dict[str, Any]) -> None:
    section = doc.sections[0]
    header = section.header
    footer = section.footer

    hp = header.paragraphs[0]
    hp.text = f"{ctx_flat.get('org.trade_name', '')}  |  {spec.get('title', '')}"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in hp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(120, 120, 120)

    fp = footer.paragraphs[0]
    classification = ctx_flat.get("doc_control.classification", spec.get("classification", "Internal"))
    fp.text = (
        f"Document ID: {spec.get('doc_id', '')}   |   Version: {ctx_flat.get('doc_control.version', '1.0')}"
        f"   |   Classification: {classification}   |   {ctx_flat.get('org.legal_name', '')}"
    )
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(120, 120, 120)


def render_docx_document(spec: Dict[str, Any], context: Dict[str, Any]) -> bytes:
    """Render a policy/procedure/plan document to DOCX bytes."""
    flat = flatten_context(context)
    brand_hex = context.get("org", {}).get("brand_color", "#1F3A5F").lstrip("#")
    r, g, b = hex_to_rgb(brand_hex)
    brand_rgb = RGBColor(r, g, b)

    doc = Document()

    # Base styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    _add_doc_control_header_footer(doc, flat, spec)

    # ---- Cover page ---------------------------------------------------------
    logo_path = context.get("org", {}).get("logo_path")
    if logo_path:
        try:
            doc.add_picture(logo_path, width=Inches(2.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    doc.add_paragraph()
    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org_p.add_run(context.get("org", {}).get("legal_name", ""))
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = brand_rgb

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title_p.add_run(spec.get("title", ""))
    trun.bold = True
    trun.font.size = Pt(26)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub_p.add_run(spec.get("standard", ""))
    srun.font.size = Pt(12)
    srun.italic = True

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mrun = meta_p.add_run(
        f"Document ID: {spec.get('doc_id', '')}\n"
        f"Version: {flat.get('doc_control.version', '1.0')}\n"
        f"Effective Date: {flat.get('doc_control.effective_date', '')}\n"
        f"Classification: {flat.get('doc_control.classification', spec.get('classification', 'Internal'))}"
    )
    mrun.font.size = Pt(11)

    doc.add_page_break()

    # ---- Document control table --------------------------------------------
    _heading(doc, "Document Control", brand_rgb, level=1)
    control_rows = [
        ("Document Title", spec.get("title", "")),
        ("Document ID", spec.get("doc_id", "")),
        ("Version", flat.get("doc_control.version", "1.0")),
        ("Standard", spec.get("standard", "")),
        ("Clause / Control Reference", ", ".join(spec.get("clause_refs", []))),
        ("Classification", flat.get("doc_control.classification", spec.get("classification", "Internal"))),
        ("Effective Date", flat.get("doc_control.effective_date", "")),
        ("Next Review Date", flat.get("doc_control.review_date", "")),
        ("Document Owner", resolve_variables(spec.get("owner_role", ""), flat)),
        ("Reviewed By", flat.get("doc_control.reviewed_by", "")),
        ("Approved By", flat.get("doc_control.approved_by", "")),
        ("Prepared By", flat.get("doc_control.prepared_by", "")),
    ]
    _kv_table(doc, control_rows, brand_hex)

    # ---- Change history -----------------------------------------------------
    doc.add_paragraph()
    _heading(doc, "Change History", brand_rgb, level=2)
    ch = doc.add_table(rows=1, cols=4)
    ch.style = "Table Grid"
    hdr = ch.rows[0].cells
    for i, label in enumerate(["Version", "Date", "Description of Change", "Author"]):
        hdr[i].text = label
        _bold_cell(hdr[i], brand_hex)
    row = ch.add_row().cells
    row[0].text = flat.get("doc_control.version", "1.0")
    row[1].text = flat.get("doc_control.effective_date", "")
    row[2].text = "Initial issue."
    row[3].text = flat.get("doc_control.prepared_by", "")

    # ---- Body sections ------------------------------------------------------
    for sec in spec.get("sections", []):
        doc.add_paragraph()
        _heading(doc, sec.get("heading", ""), brand_rgb, level=1)
        body = resolve_variables(sec.get("body", ""), flat)
        _render_body(doc, body)

        # bullet lists inside a section
        for bullet in sec.get("bullets", []) or []:
            p = doc.add_paragraph(resolve_variables(bullet, flat), style="List Bullet")
            p.paragraph_format.space_after = Pt(2)

    # ---- Roles & responsibilities table (if provided) ----------------------
    if spec.get("responsibilities"):
        doc.add_paragraph()
        _heading(doc, "Roles and Responsibilities", brand_rgb, level=1)
        rt = doc.add_table(rows=1, cols=2)
        rt.style = "Table Grid"
        h = rt.rows[0].cells
        h[0].text, h[1].text = "Role", "Responsibility"
        _bold_cell(h[0], brand_hex)
        _bold_cell(h[1], brand_hex)
        for role, resp in spec["responsibilities"]:
            rrow = rt.add_row().cells
            rrow[0].text = resolve_variables(role, flat)
            rrow[1].text = resolve_variables(resp, flat)

    # ---- Monitoring & KPIs --------------------------------------------------
    if spec.get("kpis"):
        doc.add_paragraph()
        _heading(doc, "Monitoring and KPIs", brand_rgb, level=1)
        for kpi in spec["kpis"]:
            doc.add_paragraph(resolve_variables(kpi, flat), style="List Bullet")

    # ---- Common nonconformities --------------------------------------------
    if spec.get("common_nonconformities"):
        doc.add_paragraph()
        _heading(doc, "Common Nonconformities", brand_rgb, level=1)
        for nc in spec["common_nonconformities"]:
            doc.add_paragraph(resolve_variables(nc, flat), style="List Bullet")

    # ---- Internal audit questions ------------------------------------------
    if spec.get("audit_questions"):
        doc.add_paragraph()
        _heading(doc, "Internal Audit Questions", brand_rgb, level=1)
        for q in spec["audit_questions"]:
            doc.add_paragraph(resolve_variables(q, flat), style="List Number")

    # ---- Related documents --------------------------------------------------
    if spec.get("related_documents"):
        doc.add_paragraph()
        _heading(doc, "Related Documents", brand_rgb, level=1)
        for rd in spec["related_documents"]:
            doc.add_paragraph(resolve_variables(rd, flat), style="List Bullet")

    # ---- Legal / disclaimer -------------------------------------------------
    if spec.get("disclaimer"):
        doc.add_paragraph()
        _heading(doc, "Important Notice", brand_rgb, level=2)
        dp = doc.add_paragraph()
        drun = dp.add_run(resolve_variables(spec["disclaimer"], flat))
        drun.italic = True
        drun.font.size = Pt(9)
        drun.font.color.rgb = RGBColor(150, 60, 60)

    # ---- Approval -----------------------------------------------------------
    doc.add_paragraph()
    _heading(doc, "Approval", brand_rgb, level=1)
    at = doc.add_table(rows=1, cols=4)
    at.style = "Table Grid"
    ah = at.rows[0].cells
    for i, label in enumerate(["Action", "Name", "Role", "Date"]):
        ah[i].text = label
        _bold_cell(ah[i], brand_hex)
    for action, name in [
        ("Prepared", flat.get("doc_control.prepared_by", "")),
        ("Reviewed", flat.get("doc_control.reviewed_by", "")),
        ("Approved", flat.get("doc_control.approved_by", "")),
    ]:
        arow = at.add_row().cells
        arow[0].text = action
        arow[1].text = name
        arow[2].text = ""
        arow[3].text = flat.get("doc_control.effective_date", "")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _heading(doc: Document, text: str, brand_rgb: RGBColor, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = brand_rgb
    run.font.size = Pt(15 if level == 1 else 12)


def _render_body(doc: Document, body: str) -> None:
    for para in body.split("\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)


def _kv_table(doc: Document, rows: List[Tuple[str, str]], brand_hex: str) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v
        _bold_cell(cells[0], brand_hex, white_text=True)


def _bold_cell(cell, brand_hex: str, white_text: bool = True) -> None:
    _set_cell_background(cell, brand_hex)
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            if white_text:
                run.font.color.rgb = RGBColor(255, 255, 255)


# --------------------------------------------------------------------------- #
# XLSX rendering
# --------------------------------------------------------------------------- #
def render_xlsx_register(spec: Dict[str, Any], context: Dict[str, Any]) -> bytes:
    """Render a register / tracker to XLSX bytes."""
    flat = flatten_context(context)
    brand_hex = context.get("org", {}).get("brand_color", "#1F3A5F").lstrip("#")

    wb = Workbook()
    ws = wb.active
    ws.title = (spec.get("sheet_name") or "Register")[:31]

    # Title band
    title = resolve_variables(spec.get("title", "Register"), flat)
    columns = spec.get("columns", [])
    ncols = max(len(columns), 1)

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncols)
    tcell = ws.cell(row=1, column=1, value=f"{context.get('org', {}).get('legal_name', '')} — {title}")
    tcell.font = Font(bold=True, size=14, color="FFFFFF")
    tcell.fill = PatternFill("solid", fgColor=brand_hex)
    tcell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 26

    meta_cell = ws.cell(
        row=2, column=1,
        value=(f"Document ID: {spec.get('doc_id', '')}   |   Version: "
               f"{flat.get('doc_control.version', '1.0')}   |   Standard: {spec.get('standard', '')}"),
    )
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=ncols)
    meta_cell.font = Font(italic=True, size=9, color="666666")

    # Header row
    header_row = 4
    thin = Side(style="thin", color="BBBBBB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for i, col in enumerate(columns, start=1):
        c = ws.cell(row=header_row, column=i, value=col.get("label", ""))
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor=brand_hex)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = border
        ws.column_dimensions[get_column_letter(i)].width = col.get("width", 22)

    # Seed rows
    for r_idx, row_vals in enumerate(spec.get("rows", []), start=header_row + 1):
        for c_idx, col in enumerate(columns, start=1):
            key = col.get("key", "")
            raw = row_vals.get(key, "")
            val = resolve_variables(str(raw), flat) if isinstance(raw, str) else raw
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border

    ws.freeze_panes = ws.cell(row=header_row + 1, column=1)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# PDF rendering (controlled reference copy)
# --------------------------------------------------------------------------- #
def render_pdf_document(spec: Dict[str, Any], context: Dict[str, Any], watermark: str = "") -> bytes:
    """Render a controlled PDF reference copy of a document."""
    flat = flatten_context(context)
    brand_hex = context.get("org", {}).get("brand_color", "#1F3A5F")
    r, g, b = hex_to_rgb(brand_hex)
    brand_color = colors.Color(r / 255, g / 255, b / 255)

    buf = io.BytesIO()
    doc_tmpl = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=spec.get("title", ""),
    )

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], textColor=brand_color, fontSize=15, spaceAfter=6)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=brand_color, fontSize=12)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=9.5, leading=13)
    cover_org = ParagraphStyle("CoverOrg", parent=styles["Title"], textColor=brand_color, fontSize=20, alignment=TA_CENTER)
    cover_title = ParagraphStyle("CoverTitle", parent=styles["Title"], fontSize=26, alignment=TA_CENTER)
    small = ParagraphStyle("Small", parent=styles["BodyText"], fontSize=9, textColor=colors.grey, alignment=TA_CENTER)

    flow: List[Any] = []

    # Cover
    logo_path = context.get("org", {}).get("logo_path")
    if logo_path:
        try:
            flow.append(RLImage(logo_path, width=45 * mm, height=45 * mm, kind="proportional"))
        except Exception:
            pass
    flow.append(Spacer(1, 12 * mm))
    flow.append(Paragraph(context.get("org", {}).get("legal_name", ""), cover_org))
    flow.append(Spacer(1, 6 * mm))
    flow.append(Paragraph(spec.get("title", ""), cover_title))
    flow.append(Spacer(1, 4 * mm))
    flow.append(Paragraph(spec.get("standard", ""), small))
    flow.append(Spacer(1, 10 * mm))
    flow.append(Paragraph(
        f"Document ID: {spec.get('doc_id', '')} &nbsp;|&nbsp; Version: {flat.get('doc_control.version', '1.0')}"
        f" &nbsp;|&nbsp; Classification: {flat.get('doc_control.classification', spec.get('classification', 'Internal'))}",
        small))
    flow.append(PageBreak())

    # Document control table
    flow.append(Paragraph("Document Control", h1))
    control_data = [
        ["Document Title", spec.get("title", "")],
        ["Document ID", spec.get("doc_id", "")],
        ["Version", flat.get("doc_control.version", "1.0")],
        ["Standard", spec.get("standard", "")],
        ["Clause / Control Ref.", ", ".join(spec.get("clause_refs", []))],
        ["Effective Date", flat.get("doc_control.effective_date", "")],
        ["Owner", resolve_variables(spec.get("owner_role", ""), flat)],
        ["Approved By", flat.get("doc_control.approved_by", "")],
    ]
    t = Table([[Paragraph(f"<b>{k}</b>", body), Paragraph(str(v), body)] for k, v in control_data],
              colWidths=[55 * mm, 110 * mm])
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("BACKGROUND", (0, 0), (0, -1), brand_color),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    flow.append(t)
    flow.append(Spacer(1, 6 * mm))

    # Body sections
    for sec in spec.get("sections", []):
        block = [Paragraph(sec.get("heading", ""), h1)]
        text = resolve_variables(sec.get("body", ""), flat)
        for para in text.split("\n"):
            if para.strip():
                block.append(Paragraph(para.strip().replace("&", "&amp;"), body))
        for bullet in sec.get("bullets", []) or []:
            block.append(Paragraph("• " + resolve_variables(bullet, flat).replace("&", "&amp;"), body))
        block.append(Spacer(1, 3 * mm))
        flow.append(KeepTogether(block))

    if spec.get("disclaimer"):
        disc_style = ParagraphStyle("Disc", parent=body, textColor=colors.Color(0.6, 0.2, 0.2), fontSize=8.5)
        flow.append(Spacer(1, 4 * mm))
        flow.append(Paragraph("Important Notice", h2))
        flow.append(Paragraph(resolve_variables(spec["disclaimer"], flat).replace("&", "&amp;"), disc_style))

    def _footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.grey)
        footer_text = (
            f"{spec.get('doc_id', '')}  |  v{flat.get('doc_control.version', '1.0')}  |  "
            f"{flat.get('doc_control.classification', spec.get('classification', 'Internal'))}  |  "
            f"{context.get('org', {}).get('legal_name', '')}  |  Page {_doc.page}"
        )
        canvas.drawCentredString(A4[0] / 2, 10 * mm, footer_text)
        if watermark:
            canvas.saveState()
            canvas.setFont("Helvetica-Bold", 60)
            canvas.setFillColor(colors.Color(0.9, 0.9, 0.9))
            canvas.translate(A4[0] / 2, A4[1] / 2)
            canvas.rotate(45)
            canvas.drawCentredString(0, 0, watermark)
            canvas.restoreState()
        canvas.restoreState()

    doc_tmpl.build(flow, onFirstPage=_footer, onLaterPages=_footer)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# ZIP packaging
# --------------------------------------------------------------------------- #
def build_zip(files: Dict[str, bytes]) -> bytes:
    """files: mapping of archive-path -> bytes. Returns ZIP bytes."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for path, data in files.items():
            zf.writestr(path, data)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# Validation
# --------------------------------------------------------------------------- #
def _collect_text(spec: Dict[str, Any], context: Dict[str, Any]) -> str:
    flat = flatten_context(context)
    parts: List[str] = [
        spec.get("title", ""),
        resolve_variables(spec.get("owner_role", ""), flat),
        resolve_variables(spec.get("disclaimer", ""), flat),
    ]
    for sec in spec.get("sections", []):
        parts.append(resolve_variables(sec.get("body", ""), flat))
        for b in sec.get("bullets", []) or []:
            parts.append(resolve_variables(b, flat))
    for extra_key in ("kpis", "common_nonconformities", "audit_questions", "related_documents"):
        for item in spec.get(extra_key, []) or []:
            parts.append(resolve_variables(item, flat))
    for role, resp in spec.get("responsibilities", []) or []:
        parts.append(resolve_variables(role, flat))
        parts.append(resolve_variables(resp, flat))
    return "\n".join(parts)


def validate_document(spec: Dict[str, Any], context: Dict[str, Any]) -> Dict[str, Any]:
    """Return a validation report for a single document spec + context."""
    errors: List[str] = []
    warnings: List[str] = []

    text = _collect_text(spec, context)
    unresolved = find_unresolved(text)
    if unresolved:
        errors.append(f"Unresolved placeholders: {sorted(set(unresolved))}")

    # Mandatory doc-control fields
    required_meta = ["doc_id", "title", "standard"]
    for m in required_meta:
        if not spec.get(m):
            errors.append(f"Missing required document field: {m}")

    flat = flatten_context(context)
    for dc in ["doc_control.version", "doc_control.effective_date"]:
        if not flat.get(dc):
            warnings.append(f"Missing document control value: {dc}")

    # Non-blank mandatory sections
    for sec in spec.get("sections", []):
        if sec.get("mandatory") and not (sec.get("body") or sec.get("bullets")):
            errors.append(f"Mandatory section blank: {sec.get('heading')}")

    if not context.get("org", {}).get("legal_name"):
        errors.append("Client legal name missing from context.")

    return {
        "doc_id": spec.get("doc_id"),
        "title": spec.get("title"),
        "ok": len(errors) == 0,
        "errors": errors,
        "warnings": warnings,
    }
